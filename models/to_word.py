from docx import Document
from docx.shared import Inches

# Create .docx document type
document = Document()
document.save(f'ponuda.docx')

# Create function to manipulate with that document
def export_to_word(company, names, quantity, unit, price, total, bid_num, date):
    with open (f'ponuda.docx', 'rb') as doc:
        document = Document(doc)
        name = company['Name']
        address = company['Address']
        oib = company ['OIB']
        head = document.add_table (1, 3)
        company_info = head.rows[0].cells
        company_info[0].text = f'{name}\n{address}\nOIB: {oib}'
        company_info[1].text = ''
        company_info[2].text = 'Roko d.o.o\nPožega, Dervišaga\nS.Radića 149\n34 000 Požega\nOIB:42576640228\nIBAN:HR9523860021120016238\nMob:091/22 48 88 0\nMob:091/50 44 66 3'
        document.add_paragraph ()
        document.add_paragraph (f'Mjesto i nadnevak: Dervišaga, {date}')
        document.add_paragraph (f'\t\t\t\t        Ponuda broj: {bid_num}')
        document.add_paragraph ()
        table = document.add_table (len (names) + 1, 4)
        head_cells = table.rows[0].cells
        head_cells[0].text = 'NAZIV'
        head_cells[1].text = 'MJERNA JEDINICA'
        head_cells[2].text = 'KOLIČINA'
        head_cells[3].text = 'CIJENA'
        name_cells = table.columns[0].cells
        unit_cells = table.columns[1].cells
        quantity_cells = table.columns[2].cells
        price_cells = table.columns[3].cells
        index = 1
        for cells in name_cells:
            cells.width = Inches (8)
        for i in range (len (names)):
            name_cells[index].text = names[i]
            unit_cells[index].text = unit[i]
            quantity_cells[index].text = quantity[i]
            price_cells[index].text = price[i]
            index += 1
        
        document.add_paragraph ()
        document.add_paragraph (f'\t\t\t\t\t\t\t\t\tIZNOS: {total[0]}  ')
        document.add_paragraph (f'\t\t\t\t\t\t\t\t\tPDV 25%: {total[1]}')
        document.add_paragraph (f'\t\t\t\t\t\t\t\t\tUKUPNO: {total[2]}')
        document.add_paragraph ()
        additional = document.add_table (1, 1)
        info = additional.rows[0].cells
        info_col = additional.columns[0].cells
        for cells in info_col:
            cells.width = Inches (10)
        info[0].text = 'Opcija ponude 5 dana\nRok isporuke 30 dana od primitka narudžbenice\nRobu preuzeti u skladištu ROKO d.o.o. Dervišaga'
        document.add_paragraph ()
        document.add_paragraph ('\t\t\t\t        ODGOVORNA OSOBA')
        document.add_paragraph ('\t\t\t\t\tMarijan Babić')
        document.add_paragraph ()
        document.add_paragraph (f'\t\t\t\t\t\t\t\t\tPOTPIS I PEČAT')
    # Save changes
    document.save (f'ponuda.docx')
